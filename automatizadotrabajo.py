import os
import pandas as pd
from docx import Document
from datetime import datetime

# Ruta del archivo Excel y plantilla de Word
excel_path = r'E:\PROGRAMAS JUEGOS\DOCUMENTO AUTOMATIZADO\Calificaciones.xlsx'
word_template_path = r'E:\PROGRAMAS JUEGOS\DOCUMENTO AUTOMATIZADO\Reporte_Completo_Calificaciones.docx'

# Cargar los datos del archivo Excel
try:
    df = pd.read_excel(excel_path)
    print(f"Archivo Excel cargado correctamente. Número de filas: {len(df)}")
except Exception as e:
    print(f"Error al cargar el archivo Excel: {e}")
    exit()

# Función para reemplazar las etiquetas en el documento Word (para párrafos y tablas)
def reemplazar_etiquetas(doc, nombre, matematicas, lenguaje, historia, ciencias, educacion_fisica, arte, ingles, informatica, musica, formacion_civica, promedio, fecha):
    etiquetas = {
        "{{Nombre del Alumno}}": nombre,
        "{{Matematicas}}": str(matematicas),
        "{{Lenguaje y Literatura}}": str(lenguaje),
        "{{Historia}}": str(historia),
        "{{Ciencias Naturales}}": str(ciencias),
        "{{Educacion Fisica}}": str(educacion_fisica),
        "{{Arte}}": str(arte),
        "{{Ingles}}": str(ingles),
        "{{Informatica}}": str(informatica),
        "{{Musica}}": str(musica),
        "{{Formacion Civica}}": str(formacion_civica),
        "{{Promedio}}": str(promedio),
        "{{fecha}}": fecha
    }

    # Reemplazar en cada párrafo
    for para in doc.paragraphs:
        for etiqueta, valor in etiquetas.items():
            if etiqueta in para.text:
                para.text = para.text.replace(etiqueta, valor)

    # Reemplazar en cada celda de la tabla (si existe alguna)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for etiqueta, valor in etiquetas.items():
                    if etiqueta in cell.text:
                        cell.text = cell.text.replace(etiqueta, valor)

# Obtener la fecha actual
fecha = datetime.now().strftime("%d-%m-%Y")

# Crear la subcarpeta 'Calificaciones' con la fecha
carpeta_calificaciones = f"E:\\PROGRAMAS JUEGOS\\DOCUMENTO AUTOMATIZADO\\Calificaciones_{fecha}"
if not os.path.exists(carpeta_calificaciones):
    os.makedirs(carpeta_calificaciones)

# Procesar cada estudiante en el archivo Excel
for index, row in df.iterrows():
    nombre = row['Nombre del Alumno']
    matematicas = row['Matematicas']
    lenguaje = row['Lenguaje y Literatura']
    historia = row['Historia']
    ciencias = row['Ciencias Naturales']
    educacion_fisica = row['Educacion Fisica']
    arte = row['Arte']
    ingles = row['Ingles']
    informatica = row['Informatica']
    musica = row['Musica']
    formacion_civica = row['Formacion Civica']
    promedio = row['Promedio']

    # Cargar el documento de la plantilla Word
    try:
        doc = Document(word_template_path)
    except Exception as e:
        print(f"Error al cargar el archivo Word para {nombre}: {e}")
        continue

    # Reemplazar las etiquetas con los datos del estudiante y la fecha
    reemplazar_etiquetas(doc, nombre, matematicas, lenguaje, historia, ciencias, educacion_fisica, arte, ingles, informatica, musica, formacion_civica, promedio, fecha)

    # Guardar el documento en la subcarpeta 'Calificaciones' con la fecha
    try:
        ruta_guardado = os.path.join(carpeta_calificaciones, f"{nombre}_calificaciones_{fecha}.docx")
        doc.save(ruta_guardado)
        print(f"Documento generado para {nombre} en la ruta: {ruta_guardado}")
    except Exception as e:
        print(f"Error al guardar el documento para {nombre}: {e}")
